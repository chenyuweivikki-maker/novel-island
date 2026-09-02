"""
多模态文档解析 — PRD 场景二 / P2-5：word、pdf、文本 多模态上传

支持：
  .txt / .md   — 直接读文本
  .docx        — python-docx 提取段落+表格
  .pdf         — pypdf 逐页提取文本

失败时有明确错误信息（缺依赖/加密 PDF/空文件）。
"""
import io
from typing import Optional


def parse_text_file(content: bytes, filename: str) -> str:
    """按扩展名分发解析，返回提取的文本"""
    name = (filename or "").lower()
    if name.endswith((".txt", ".md", ".markdown")):
        return content.decode("utf-8", errors="replace")
    if name.endswith(".docx"):
        return _parse_docx(content)
    if name.endswith(".pdf"):
        return _parse_pdf(content)
    if name.endswith((".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp")):
        return _ocr_image(content)
    raise ValueError(f"不支持的文件类型：{filename or '未知'}（支持 .txt / .md / .docx / .pdf / 图片(png/jpg)）")


def _parse_docx(content: bytes) -> str:
    """python-docx 提取段落与表格文本"""
    try:
        from docx import Document
    except ImportError:
        raise ValueError("docx 解析依赖未安装（pip install python-docx），请先安装或改用 .txt/.md")
    doc = Document(io.BytesIO(content))
    parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts)
    if not text.strip():
        raise ValueError("docx 中没有可提取的文本")
    return text


def _parse_pdf(content: bytes) -> str:
    """pypdf 逐页提取文本（pypdf 已在 venv）"""
    try:
        import pypdf
    except ImportError:
        raise ValueError("pdf 解析依赖未安装（pip install pypdf），请先安装或改用 .txt/.md")
    try:
        reader = pypdf.PdfReader(io.BytesIO(content))
    except Exception as e:
        raise ValueError(f"PDF 解析失败（可能已加密或损坏）：{e}")
    parts = []
    for page in reader.pages:
        try:
            t = page.extract_text() or ""
        except Exception:
            t = ""
        if t.strip():
            parts.append(t)
    text = "\n".join(parts)
    if not text.strip():
        raise ValueError("PDF 中没有可提取的文本（可能是扫描件，OCR 二期支持）")
    return text


def _ocr_image(content: bytes) -> str:
    """图片 OCR（P2-5）：用 macOS Vision 框架识别图片文字，返回文本。

    仅 macOS 可用（依赖 pyobjc-framework-Vision）。零成本、本地识别、无需 API Key。
    """
    try:
        import Vision
        from Foundation import NSData
    except ImportError:
        raise ValueError("图片 OCR 需 macOS Vision（python -m pip install pyobjc-framework-Vision）")
    data = NSData.dataWithBytes_length_(content, len(content))
    req = Vision.VNRecognizeTextRequest.alloc().init()
    req.setRecognitionLanguages_(["zh-Hans", "en-US"])
    req.setRecognitionLevel_(Vision.VNRequestTextRecognitionLevelAccurate)
    handler = Vision.VNImageRequestHandler.alloc().initWithData_options_(data, {})
    ok, err = handler.performRequests_error_([req], None)
    if not ok:
        raise ValueError(f"OCR 识别失败：{err}")
    lines = []
    for obs in req.results() or []:
        cands = obs.topCandidates_(1)
        if cands and len(cands):
            lines.append(str(cands[0].string()))
    text = "\n".join(lines)
    if not text.strip():
        raise ValueError("图片中没有识别到文字")
    return text
